"""
File handlers for different document types.
"""
import json
import io
from typing import Dict, Any, Optional
from PyPDF2 import PdfReader
from docx import Document


class BaseHandler:
    """Base class for file handlers."""
    
    def extract_content(self, file_data: bytes) -> Dict[str, Any]:
        """Extract content from file data."""
        raise NotImplementedError


class PDFHandler(BaseHandler):
    """Handler for PDF files."""
    
    def extract_content(self, file_data: bytes) -> Dict[str, Any]:
        """Extract text content from PDF."""
        try:
            pdf_file = io.BytesIO(file_data)
            reader = PdfReader(pdf_file)
            
            text_content = []
            num_pages = len(reader.pages)
            
            for page_num, page in enumerate(reader.pages[:10]):  # Limit to first 10 pages for preview
                text = page.extract_text()
                text_content.append(text)
            
            full_text = "\n".join(text_content)
            
            # Extract metadata if available
            metadata = {}
            if reader.metadata:
                metadata = {
                    "title": reader.metadata.get("/Title", ""),
                    "author": reader.metadata.get("/Author", ""),
                    "subject": reader.metadata.get("/Subject", ""),
                    "creator": reader.metadata.get("/Creator", ""),
                    "producer": reader.metadata.get("/Producer", ""),
                    "num_pages": num_pages
                }
            
            return {
                "extracted_text": full_text,
                "content_preview": full_text[:500],  # First 500 characters
                "metadata": metadata
            }
        except Exception as e:
            return {
                "extracted_text": "",
                "content_preview": f"Error extracting PDF content: {str(e)}",
                "metadata": {}
            }


class JSONHandler(BaseHandler):
    """Handler for JSON files."""
    
    def extract_content(self, file_data: bytes) -> Dict[str, Any]:
        """Extract content from JSON file."""
        try:
            # Try to decode as UTF-8
            text = file_data.decode("utf-8")
            json_data = json.loads(text)
            
            # Convert JSON to readable string
            json_string = json.dumps(json_data, indent=2)
            
            # Extract metadata from JSON structure
            metadata = {}
            if isinstance(json_data, dict):
                # Look for common metadata fields
                for key in ["title", "name", "id", "type", "version", "author", "created", "modified"]:
                    if key in json_data:
                        metadata[key] = json_data[key]
            
            return {
                "extracted_text": json_string,
                "content_preview": json_string[:500],
                "metadata": metadata,
                "json_data": json_data  # Store parsed JSON for easier querying
            }
        except UnicodeDecodeError:
            return {
                "extracted_text": "",
                "content_preview": "Error: File is not valid UTF-8 encoded JSON",
                "metadata": {}
            }
        except json.JSONDecodeError as e:
            return {
                "extracted_text": "",
                "content_preview": f"Error parsing JSON: {str(e)}",
                "metadata": {}
            }


class WordHandler(BaseHandler):
    """Handler for Word documents (.docx)."""
    
    def extract_content(self, file_data: bytes) -> Dict[str, Any]:
        """Extract text content from Word document."""
        try:
            doc_file = io.BytesIO(file_data)
            doc = Document(doc_file)
            
            # Extract text from all paragraphs
            paragraphs = [para.text for para in doc.paragraphs]
            full_text = "\n".join(paragraphs)
            
            # Extract text from tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    table_texts.append(row_text)
            
            if table_texts:
                full_text += "\n\nTables:\n" + "\n".join(table_texts)
            
            # Extract metadata
            metadata = {}
            core_props = doc.core_properties
            if core_props:
                metadata = {
                    "title": core_props.title or "",
                    "author": core_props.author or "",
                    "subject": core_props.subject or "",
                    "keywords": core_props.keywords or "",
                    "created": str(core_props.created) if core_props.created else "",
                    "modified": str(core_props.modified) if core_props.modified else "",
                    "num_paragraphs": len(paragraphs),
                    "num_tables": len(doc.tables)
                }
            
            return {
                "extracted_text": full_text,
                "content_preview": full_text[:500],
                "metadata": metadata
            }
        except Exception as e:
            return {
                "extracted_text": "",
                "content_preview": f"Error extracting Word document content: {str(e)}",
                "metadata": {}
            }

